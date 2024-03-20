import json
from dataclasses import dataclass, field, is_dataclass, asdict
from typing import List, Optional
import re

import docx
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx import Document
from docx.text.paragraph import Paragraph

from common.logger import logger

from utils import find_text_under_header
from common.utils import get_path

NOT_FORMULATED_YET = "Not formulated yet."
NOT_AVAILABLE_THIS_IS_ROOT_SECTION = "Not available. This is the root section."


@dataclass
class Section:
    title: str
    abstract: str
    content: str
    sources: List[str] = field(default_factory=list)
    feedback: str = ""
    subsections: List["Section"] = field(default_factory=list)
    parent: Optional["Section"] = None

    def render_table_of_contents(self, include_abstract: bool = True) -> str:
        # Recursively renders the table of contents
        message = f"## {self.title}\n\n"
        if include_abstract:
            message += f"{self.abstract}\n\n"
        for subsection in self.subsections:
            message += subsection.render_table_of_contents()
        return message

    def render(self) -> str:
        # Renders everything as a kind of key-value store and just slaps that into the user prompt
        message = f"# Title\n\n{self.title}\n\n## Abstract\n\n{self.abstract}\n\n"
        message += f"## Content\n\n{self.content}\n\n"
        if self.feedback:
            message += f"## Feedback\n\n{self.feedback}\n\n"
        subsections = ""
        for subsection in self.subsections:
            subsections += subsection.render_as_subsection()
        message += f"## Subsections\n\n{subsections}"
        message += f"## Sources\n\n - " + "\n - ".join(self.sources)
        return message

    def render_as_text(self, header_level: int = 1, show_abstracts: bool = True) -> str:
        header_hash = "#" * header_level
        if show_abstracts:
            rendered = f"{header_hash} {self.title.strip()}\n\n**Abstract**: {self.abstract.strip()}\n\n{self.content.strip()}\n\n"
        else:
            rendered = f"{header_hash} {self.title}\n\n{self.content}"
        for subsection in self.subsections:
            rendered += subsection.render_as_text(
                header_level=header_level + 1, show_abstracts=show_abstracts
            )
        return rendered

    def render_full_tree(self, current_depth: int = 0) -> str:
        # renders the whole document tree, where section markers are just the number of #s corresponding to the depth of the section. Note the # Subsections there at the end, which I put in to try to emphasive the hierarachy of the document... could cut that?
        current_depth += 1
        section_marker = "#" * current_depth
        message = f"{section_marker} {self.title}\n\n{self.abstract}\n\n"
        subsections = ""
        for subsection in self.subsections:
            subsections += subsection.render_full_tree(current_depth)
        if subsections:
            message += f"{section_marker}# Subsections\n\n{subsections}"
        return message

    def render_sources(self) -> str:
        return " - " + "\n - ".join(self.sources) + "\n\n"

    def render_as_subsection(self, header_level: int = 1) -> str:
        base_hash = "#" * header_level
        sub_hash = "#" * (header_level + 1)
        return f"{base_hash} Subsection Title\n{self.title}\n\n{sub_hash} Subsection Abstract\n{self.abstract}\n\n"

    def render_title_and_abstract(self) -> str:
        return f"# Title\n\n{self.title}\n\n## Abstract\n\n{self.abstract}"

    def render_subsections(self, header_level: int = 1) -> str:
        subsections = ""
        for subsection in self.subsections:
            subsections += subsection.render_as_subsection(header_level=header_level)
        return subsections

    def clone_and_update_from_string(self, string):
        new_section = self.__class__.from_string(string)
        new_section.subsections = self.subsections
        return new_section

    @classmethod
    def from_string(cls, string: str, parse_subsections: bool = True) -> "Section":
        # TODO : make this suck less, it is wildly unreliably and expensive since it basically has to happen twice.

        title = find_text_under_header(string, header="## Title", assert_found=True)
        abstract = find_text_under_header(
            string, header="## Abstract", assert_found=True
        )
        content = find_text_under_header(string, header="## Content", assert_found=True)

        sources = find_text_under_header(
            string, header="## Sources", assert_found=False
        )
        if sources is not None:
            sources = [
                source.replace("- ", "") for source in sources.strip().split("\n")
            ]

        subsections_string = find_text_under_header(
            string, header="## Subsections", assert_found=False
        )
        subsections = []

        if subsections_string is not None and parse_subsections:
            lines = subsections_string.split("\n")

            if len(lines) % 2 != 0:
                logger.error(
                    f"Subsections must be in pairs of title and abstract, but got {len(lines)} lines."
                )
                lines = lines[:-1]
            for line_idx, line in enumerate(lines):
                # Every even line is a title, every odd line is an abstract
                if line_idx % 2 == 0:
                    subsection_title = line.replace("-", " ").strip()
                    subsection_abstract = lines[line_idx + 1].replace("-", " ").strip()
                    subsections.append(
                        cls(
                            title=subsection_title,
                            abstract=subsection_abstract,
                            content=NOT_FORMULATED_YET,
                            sources=[],
                        )
                    )
        if title is not None:
            title = title.strip()
        if abstract is not None:
            abstract = abstract.strip()
        if content is not None:
            content = content.strip()

        return cls(
            title=title,
            abstract=abstract,
            content=content,
            sources=sources,
            subsections=subsections,
        )

    def to_dict(self):
        if is_dataclass(self):
            return asdict(self)
        return self

    @classmethod
    def from_dict(cls, dict_obj):
        if "subsections" in dict_obj:
            dict_obj["subsections"] = [
                cls.from_dict(subsection) for subsection in dict_obj["subsections"]
            ]
        if "parent" in dict_obj and dict_obj["parent"] is not None:
            dict_obj["parent"] = cls.from_dict(dict_obj["parent"])
        if "sources" not in dict_obj:
            dict_obj["sources"] = []
        return cls(**dict_obj)

    def persist(self, path: str) -> None:
        with open(path, "w") as f:
            json.dump(self.to_dict(), f)

    @classmethod
    def load(cls, path: str) -> "Section":
        with open(path, "r") as f:
            return cls.from_dict(json.load(f))

    def export_to_docx(self, path: Optional[str] = None) -> Document:
        return convert_to_docx(self, path)


def add_styled_paragraph(doc: Document, text: str, style: str = None, size: int = None):
    if size is not None:
        size = Pt(size)

    p = doc.add_paragraph(style=style)
    bold_re = r"\*\*(.*?)\*\*|__(.*?)__"
    italic_re = r"\*(.*?)\*|_(.*?)_"
    link_re = r"\[([^\]]+)\]\(([^)]+)\)"

    while text:
        bold_match = re.search(bold_re, text)
        italic_match = re.search(italic_re, text)
        link_match = re.search(link_re, text)

        nearest_match = None
        for match in [bold_match, italic_match, link_match]:
            if match and (
                nearest_match is None or match.start() < nearest_match.start()
            ):
                nearest_match = match

        if nearest_match:
            run = p.add_run(text[: nearest_match.start()])

            if size:
                run.font.size = size

            if nearest_match is bold_match:
                styled_text = nearest_match.group(1) or nearest_match.group(2)
                run = p.add_run(styled_text)
                run.bold = True
            elif nearest_match is italic_match:
                styled_text = nearest_match.group(1) or nearest_match.group(2)
                run = p.add_run(styled_text)
                run.italic = True
            elif nearest_match is link_match:
                link_text = nearest_match.group(1)
                link_url = nearest_match.group(2)
                run = add_hyperlink(p, link_text, link_url)

            if size:
                run.font.size = size

            text = text[nearest_match.end() :]
        else:
            run = p.add_run(text)
            if size:
                run.font.size = size
            break


def add_hyperlink(paragraph: Paragraph, text: str, url: str):
    """
    Add a hyperlink to a paragraph with a specified color.
    """
    # This is a workaround to add hyperlinks in python-docx
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run element for hyperlink text
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set color for hyperlink (e.g., blue)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")  # Royal blue color in hex
    rPr.append(color)

    # Optionally, underline the hyperlink text
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Append properties and text to the run element
    new_run.append(rPr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    # Append the run element to the hyperlink and then to the paragraph
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def replace_citations(line: str, source_offset: int):
    # Matches [1], [^1], and [^1^]
    pattern = r"\[\^?(\d+)\^?\]"

    def repl(match):
        # Extract the number from the match, add the source offset, and return the new citation
        num = int(match.group(1))
        return f"[{source_offset + num}]"

    # Replace all citations in the line using the repl function
    line = re.sub(pattern, repl, line)
    return line


def process_markdown(doc: Document, md_text: str, source_offset: int = 0) -> None:
    lines = md_text.split("\n")
    bullet_list_items = []
    for line in lines:
        line = replace_citations(line, source_offset)

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif re.match(r"^\s*[\*\+\-]\s+", line):
            line = re.sub(r"^\s*[\*\+\-]\s+", "", line)
            bullet_list_items.append(line)
        elif line.strip() != "":
            if bullet_list_items:
                add_bullet_list(doc, bullet_list_items)
                bullet_list_items = []
            add_styled_paragraph(doc, line)

    # Add any remaining bullet list items
    if bullet_list_items:
        add_bullet_list(doc, bullet_list_items)


def add_bullet_list(doc: Document, items: List[str]) -> None:
    for item in items:
        add_styled_paragraph(doc, item, style="List Bullet 2")


def add_numbered_list(doc: Document, items: List[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet 2")


def add_abstract(doc: Document, text: str):
    # Assuming you want to use Pt 10 for the abstract
    add_styled_paragraph(doc, f"**Summary:** {text}\n", size=10)


def replace_citations_with_links(text: str, sources: List[str]) -> str:
    # Matches [1], [^1], and [^1^]
    pattern = r"\[\^?(\d+)\^?\]"

    def get_last_parentheses_content(s):
        matches = re.findall(r"\((.*?)\)", s)
        if matches:
            return matches[-1]  # Return the last match

    def repl(match):
        # Extract the number from the match and return the new citation with the link
        num = int(match.group(1))
        if (
            0 < num <= len(sources)
        ):  # Ensure num is within the bounds of the sources list
            return f"[[{num}]({get_last_parentheses_content(sources[num-1])})]"
        else:
            # FIXME: This is a temporary fix to handle out-of-bounds citation numbers
            # Optionally, handle the out-of-bounds citation number in some way
            # For now, just return the original match (citation marker) unmodified
            return match.group(0)

    # Replace all citations in the text using the repl function
    text = re.sub(pattern, repl, text)
    return text


def process_section(
    doc: Document,
    section: Section,
    sources: List[str],
    *,
    level: int = 1,
    source_offset: int = 0,
) -> int:
    doc.add_heading(section.title, level=level)

    # Identify which sources are used in the content
    used_sources = set(
        int(match.group(1))
        for match in re.finditer(r"\[\^?(\d+)\^?\]", section.content)
    )

    # Create a mapping from old citation numbers to new citation numbers
    citation_mapping = {old: new for new, old in enumerate(used_sources, 1)}

    # Replace the old citation numbers with the new citation numbers in the content
    section.content = re.sub(
        r"\[\^?(\d+)\^?\]",
        lambda match: f"[{citation_mapping[int(match.group(1))]}]",
        section.content,
    )

    # Create a new list of sources where each source is associated with its new index
    new_sources = [
        re.sub(r"^\[\d+\]", f"[{citation_mapping[i]}]", source)
        for i, source in enumerate(section.sources, 1)
        if i in used_sources
    ]
    section.sources = new_sources

    # Replace citations like [1] in section.content with citations like [1](https://example.com) using the sources
    content = replace_citations_with_links(section.content, section.sources)

    # Update the source index to include the offset for the current section
    process_markdown(doc, content, source_offset=source_offset)

    # Update the source index to include the offset for the current section
    updated_sources = [
        replace_citations(source, source_offset) for source in section.sources
    ]
    sources.extend(updated_sources)
    source_offset += len(section.sources)

    for subsection in section.subsections:
        source_offset = process_section(
            doc,
            subsection,
            level=level + 1,
            sources=sources,
            source_offset=source_offset,
        )
    return source_offset


def add_references(doc: Document, sources: List[str]) -> None:
    # Add a heading for the references section
    doc.add_heading("References", level=1)

    # Iterate through the sources and add each as a new paragraph
    for i, source in enumerate(sources, start=1):
        add_styled_paragraph(doc, f"{source}")


def convert_to_docx(section: Section, docx_file: Optional[str] = None) -> Document:
    doc = Document()

    sources = []
    sections = [section]

    # Add a header to the document
    header = doc.sections[0].header
    paragraph = header.add_paragraph()  # Create a new paragraph
    # Add your company logo to the header
    run = paragraph.add_run()
    run.add_picture(get_path("app/static/assets/tiptree-header.png"), width=Pt(300))
    run.add_break()

    for section_ in sections:
        process_section(doc, section_, sources, level=0)
    if sources:
        add_references(doc, sources)
    if docx_file is not None:
        doc.save(docx_file)
    return doc
